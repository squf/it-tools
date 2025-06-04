import os
import io
from docx import Document
from docx.shared import Inches

# Function to extract the first image or the one after "Delegation" -- need to update this if statement to match your use case, in my case i wanted things after the word "Delegation" appeared in each .docx file!
def extract_relevant_image(doc):
    found_delegation = False
    for para in doc.paragraphs:
        if "Delegation" in para.text:
            found_delegation = True
            break

    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            return rel.target_part.blob  # Return the first image blob -- you would need to tweak this as well if you're trying to target something besides the first screenshot in the .docx file
    return None

# Process a single .docx file
def process_docx_file(file_path):
    doc = Document(file_path)
    file_name = os.path.basename(file_path)
    image_blob = extract_relevant_image(doc)
    return file_name, image_blob

# Process all .docx files in a folder
def process_all_docx_files(folder_path):
    data = []
    for file_name in os.listdir(folder_path):
        if file_name.endswith(".docx"):
            file_path = os.path.join(folder_path, file_name)
            data.append(process_docx_file(file_path))
    return data

# Merge into a single .docx file
def merge_docx_files(data, output_file):
    merged_doc = Document()
    for file_name, image_blob in data:
        merged_doc.add_heading(file_name, level=1)
        if image_blob:
            image_stream = io.BytesIO(image_blob)
            merged_doc.add_picture(image_stream, width=Inches(5.5))
    merged_doc.save(output_file)

# Main function -- update folder_path to point to your folder where all of your .docx files are
def main():
    folder_path = '/home/path/to/your/files'
    output_file = 'Condensed_GPO_Docs.docx'
    data = process_all_docx_files(folder_path)
    merge_docx_files(data, output_file)

if __name__ == "__main__":
    main()
