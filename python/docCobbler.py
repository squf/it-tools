# ya this is generated code big whup

import os
import io
import pytesseract
from PIL import Image
from docx import Document
import pandas as pd
import openpyxl

# Function to extract text from image using OCR
def extract_text_from_image(image_blob):
    image_stream = io.BytesIO(image_blob)
    return pytesseract.image_to_string(Image.open(image_stream))

# Function to process a single .docx file
def process_docx_file(file_path):
    doc = Document(file_path)
    file_name = os.path.basename(file_path)
    extracted_text = ""
    for para in doc.paragraphs:
        if para.text.startswith("Links:"):
            # Find the image in the document relationships
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_blob = rel.target_part.blob
                    extracted_text = extract_text_from_image(image_blob)
                    break
            break
    return file_name, extracted_text

# Function to process all .docx files in a folder
def process_all_docx_files(folder_path):
    data = []
    for file_name in os.listdir(folder_path):
        if file_name.endswith(".docx"):
            file_path = os.path.join(folder_path, file_name)
            data.append(process_docx_file(file_path))
    return data

# Function to save data to Excel
def save_to_excel(data, output_file):
    df = pd.DataFrame(data, columns=["File Name", "Extracted Text"])
    df.to_excel(output_file, index=False)

# Function to merge all .docx files into one
def merge_docx_files(folder_path, output_file):
    merged_doc = Document()
    for file_name in os.listdir(folder_path):
        if file_name.endswith(".docx"):
            file_path = os.path.join(folder_path, file_name)
            doc = Document(file_path)
            merged_doc.add_heading(file_name, level=1)
            for para in doc.paragraphs:
                merged_doc.add_paragraph(para.text)
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_blob = rel.target_part.blob
                    image_stream = io.BytesIO(image_blob)
                    merged_doc.add_picture(image_stream)
    merged_doc.save(output_file)

# Main function - update folder_path to point to where you have saved all of your .docx files in a folder in your linux environment
def main():
    folder_path = '/home/your/path/here'
    excel_output_file = 'GPO_Docs_Extracted.xlsx'
    merged_docx_output_file = 'Merged_GPO_Docs.docx'

    # Process all .docx files and save to Excel
    data = process_all_docx_files(folder_path)
    save_to_excel(data, excel_output_file)

    # Merge all .docx files into one
    merge_docx_files(folder_path, merged_docx_output_file)

if __name__ == "__main__":
    main()
